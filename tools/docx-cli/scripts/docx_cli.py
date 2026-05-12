#!/usr/bin/env python3
"""Docx CLI - Command-line interface for python-docx.

Wraps python-docx functionality as CLI commands so agents can read, create,
and modify .docx files without writing Python code.
"""

import json
import sys
from pathlib import Path

import click
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK_TYPE
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _load(filepath):
    """Load a document. Create if it doesn't exist."""
    path = Path(filepath)
    if path.exists():
        return Document(str(path))
    raise click.ClickException(f"File not found: {filepath}")


def _save(doc, filepath, output):
    """Save document to output if given, else overwrite filepath."""
    target = output if output else filepath
    doc.save(target)


def _resolve_output(filepath, output):
    """Return (target for reading, output for saving, is_inplace)."""
    if output:
        return filepath, output, False
    return filepath, filepath, True


def _json_out(data, flag):
    """Print data as JSON if flag is set, else return None."""
    if flag:
        click.echo(json.dumps(data, ensure_ascii=False, indent=2))
    return flag


def _parse_json(ctx, param, value):
    """Parse a JSON string option."""
    if value is None:
        return None
    try:
        return json.loads(value)
    except json.JSONDecodeError as e:
        raise click.BadParameter(f"Invalid JSON: {e}")


def _parse_alignment(value):
    """Convert alignment string to WD_ALIGN_PARAGRAPH enum."""
    if value is None:
        return None
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(value.lower())


def _parse_line_spacing(value):
    """Convert line spacing string to WD_LINE_SPACING enum."""
    if value is None:
        return None
    mapping = {
        "single": WD_LINE_SPACING.SINGLE,
        "1.5": WD_LINE_SPACING.ONE_POINT_FIVE,
        "double": WD_LINE_SPACING.DOUBLE,
    }
    return mapping.get(value.lower())


def _parse_color(value):
    """Parse color from hex string or name to RGBColor."""
    if value is None:
        return None
    if value.startswith("#"):
        value = value.lstrip("#")
    try:
        r, g, b = bytes.fromhex(value.ljust(6, "0")[:6])
        return RGBColor(r, g, b)
    except (ValueError, TypeError):
        raise click.BadParameter(f"Invalid color: {value}. Use hex format (RRGGBB or #RRGGBB).")


def _parse_bool_tristate(value):
    """Parse tri-state bool option (None means not provided)."""
    return value


def _format_bool(val):
    return "yes" if val else "no"


def _format_pformat(pf):
    """Format paragraph format info."""
    info = {}
    if pf.alignment is not None:
        info["alignment"] = str(pf.alignment)
    if pf.line_spacing is not None:
        info["line_spacing"] = pf.line_spacing
    if pf.line_spacing_rule is not None:
        info["line_spacing_rule"] = str(pf.line_spacing_rule)
    if pf.space_before is not None:
        info["space_before"] = pf.space_before.pt if pf.space_before else None
    if pf.space_after is not None:
        info["space_after"] = pf.space_after.pt if pf.space_after else None
    if pf.first_line_indent is not None:
        info["first_line_indent"] = pf.first_line_indent.pt if pf.first_line_indent else None
    if pf.left_indent is not None:
        info["left_indent"] = pf.left_indent.pt if pf.left_indent else None
    if pf.right_indent is not None:
        info["right_indent"] = pf.right_indent.pt if pf.right_indent else None
    return info


def _format_font(font):
    """Format font info."""
    info = {}
    if font.name:
        info["name"] = font.name
    if font.size:
        info["size"] = font.size.pt
    if font.bold is not None:
        info["bold"] = font.bold
    if font.italic is not None:
        info["italic"] = font.italic
    if font.underline is not None:
        info["underline"] = font.underline
    if font.color and font.color.rgb:
        info["color"] = str(font.color.rgb)
    if font.highlight_color is not None:
        info["highlight"] = str(font.highlight_color)
    if font.subscript:
        info["subscript"] = True
    if font.superscript:
        info["superscript"] = True
    if font.strike:
        info["strike"] = True
    if font.double_strike:
        info["double_strike"] = True
    if font.all_caps:
        info["all_caps"] = True
    if font.small_caps:
        info["small_caps"] = True
    if font.hidden:
        info["hidden"] = True
    return info


def _set_run_font(run, bold, italic, underline, size, font_name, color,
                  highlight, superscript, subscript, strike, double_strike,
                  all_caps, small_caps, hidden):
    """Apply font formatting options to a run. Only apply non-None values."""
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline
    if size is not None:
        run.font.size = Pt(size)
    if font_name is not None:
        run.font.name = font_name
    if color is not None:
        run.font.color.rgb = _parse_color(color)
    if highlight is not None:
        from docx.enum.text import WD_COLOR_INDEX
        try:
            run.font.highlight_color = getattr(WD_COLOR_INDEX, highlight.upper())
        except AttributeError:
            raise click.BadParameter(f"Invalid highlight color: {highlight}")
    if superscript is not None:
        run.font.superscript = superscript
    if subscript is not None:
        run.font.subscript = subscript
    if strike is not None:
        run.font.strike = strike
    if double_strike is not None:
        run.font.double_strike = double_strike
    if all_caps is not None:
        run.font.all_caps = all_caps
    if small_caps is not None:
        run.font.small_caps = small_caps
    if hidden is not None:
        run.font.hidden = hidden


# ---------------------------------------------------------------------------
# CLI Entry Point
# ---------------------------------------------------------------------------

@click.group()
def cli():
    """Manipulate .docx files from the command line."""


# ---------------------------------------------------------------------------
# read
# ---------------------------------------------------------------------------

@cli.command()
@click.argument("file")
@click.option("--paragraphs/--no-paragraphs", default=True, help="Include paragraphs")
@click.option("--tables/--no-tables", default=True, help="Include tables")
@click.option("--sections/--no-sections", default=False, help="Include sections")
@click.option("--styles/--no-styles", default=False, help="Include styles")
@click.option("--headers-footers/--no-headers-footers", default=False, help="Include headers/footers")
@click.option("--images/--no-images", default=False, help="Include inline images")
@click.option("--properties/--no-properties", default=False, help="Include doc properties")
@click.option("--json", "output_json", is_flag=True, help="Output as JSON")
def read_cmd(file, paragraphs, tables, sections, styles, headers_footers, images, properties, output_json):
    """Read and display content of a .docx file."""
    doc = _load(file)
    result = {}

    # --- paragraphs ---
    if paragraphs:
        paras = []
        for i, p in enumerate(doc.paragraphs):
            pdata = {"index": i, "text": p.text}
            if p.style:
                pdata["style"] = p.style.name
            pf = _format_pformat(p.paragraph_format)
            if pf:
                pdata["format"] = pf
            # runs detail
            runs = []
            for j, r in enumerate(p.runs):
                rdata = {"index": j, "text": r.text}
                finfo = _format_font(r.font)
                if finfo:
                    rdata["font"] = finfo
                if r.style:
                    rdata["style"] = r.style.name
                runs.append(rdata)
            if runs:
                pdata["runs"] = runs
            paras.append(pdata)
        result["paragraphs"] = paras
        if not output_json:
            _pprint_paragraphs(paras)

    # --- tables ---
    if tables:
        tbls = []
        for i, t in enumerate(doc.tables):
            rows_data = []
            for ri, row in enumerate(t.rows):
                cells = [{"col": ci, "text": cell.text} for ci, cell in enumerate(row.cells)]
                rows_data.append({"index": ri, "cells": cells})
            tdata = {
                "index": i,
                "row_count": len(t.rows),
                "col_count": len(t.columns),
                "rows": rows_data,
            }
            if t.style:
                tdata["style"] = t.style.name
            tbls.append(tdata)
        result["tables"] = tbls
        if not output_json:
            _pprint_tables(tbls)

    # --- sections ---
    if sections:
        secs = []
        for i, s in enumerate(doc.sections):
            sdata = {
                "index": i,
                "page_width": s.page_width.pt if s.page_width else None,
                "page_height": s.page_height.pt if s.page_height else None,
                "orientation": str(s.orientation),
                "start_type": str(s.start_type),
                "left_margin": s.left_margin.pt if s.left_margin else None,
                "right_margin": s.right_margin.pt if s.right_margin else None,
                "top_margin": s.top_margin.pt if s.top_margin else None,
                "bottom_margin": s.bottom_margin.pt if s.bottom_margin else None,
            }
            secs.append(sdata)
        result["sections"] = secs
        if not output_json:
            _pprint_sections(secs)

    # --- styles ---
    if styles:
        stls = []
        for s in doc.styles:
            stls.append({"name": s.name, "style_id": s.style_id, "type": str(s.type)})
        result["styles"] = stls
        if not output_json:
            _pprint_styles(stls)

    # --- headers/footers ---
    if headers_footers:
        hf_data = []
        for i, s in enumerate(doc.sections):
            for hf_type in ["header", "footer", "first_page_header", "first_page_footer"]:
                try:
                    hf_obj = getattr(s, hf_type, None)
                    if hf_obj and hf_obj.paragraphs:
                        text = " | ".join(p.text for p in hf_obj.paragraphs if p.text.strip())
                        if text:
                            hf_data.append({"section": i, "type": hf_type, "text": text})
                except Exception:
                    pass
        result["headers_footers"] = hf_data
        if not output_json:
            _pprint_headers_footers(hf_data)

    # --- images ---
    if images:
        imgs = []
        for i, p in enumerate(doc.paragraphs):
            for j, r in enumerate(p.runs):
                for s in r._element.findall(qn("w:drawing")):
                    imgs.append({"paragraph": i, "run": j, "type": "drawing"})
                for s in r._element.findall(qn("w:pict")):
                    imgs.append({"paragraph": i, "run": j, "type": "pict"})
        for s in doc.inline_shapes:
            imgs.append({"type": "inline_shape", "width": s.width, "height": s.height})
        result["images"] = imgs
        if not output_json:
            _pprint_images(imgs)

    # --- properties ---
    if properties:
        cp = doc.core_properties
        props = {
            "author": cp.author,
            "title": cp.title,
            "subject": cp.subject,
            "keywords": cp.keywords,
            "category": cp.category,
            "comments": cp.comments,
            "created": str(cp.created) if cp.created else None,
            "modified": str(cp.modified) if cp.modified else None,
            "last_modified_by": cp.last_modified_by,
            "revision": cp.revision,
        }
        result["properties"] = props
        if not output_json:
            _pprint_properties(props)

    if output_json:
        click.echo(json.dumps(result, ensure_ascii=False, indent=2))


def _pprint_paragraphs(paras):
    click.echo(click.style("=== Paragraphs ===", bold=True))
    for p in paras:
        style_str = f" [{p['style']}]" if p.get("style") else ""
        click.echo(f"[{p['index']}]{style_str} {p['text']}")
        for r in p.get("runs", []):
            finfo = ""
            if r.get("font"):
                f = r["font"]
                parts = []
                if f.get("bold"):
                    parts.append("bold")
                if f.get("italic"):
                    parts.append("italic")
                if f.get("name"):
                    parts.append(f["name"])
                if f.get("size"):
                    parts.append(f"{f['size']}pt")
                if f.get("color"):
                    parts.append(f"#{f['color']}")
                if parts:
                    finfo = f" ({', '.join(parts)})"
            click.echo(f"  run[{r['index']}]{finfo}: {r['text']}")


def _pprint_tables(tbls):
    if not tbls:
        return
    click.echo(click.style("\n=== Tables ===", bold=True))
    for t in tbls:
        style_str = f" [{t['style']}]" if t.get("style") else ""
        rc = t.get("row_count", len(t["rows"]))
        cc = t.get("col_count", len(t["rows"][0]["cells"]) if t["rows"] else 0)
        click.echo(f"Table[{t['index']}]: {rc}x{cc}{style_str}")
        for row in t["rows"]:
            cells = " | ".join(c["text"] for c in row["cells"])
            click.echo(f"  [{row['index']}] {cells}")


def _pprint_sections(secs):
    if not secs:
        return
    click.echo(click.style("\n=== Sections ===", bold=True))
    for s in secs:
        click.echo(f"Section[{s['index']}]: "
                   f"{s['page_width']}x{s['page_height']}pt, "
                   f"{s['orientation']}, start={s['start_type']}")


def _pprint_styles(stls):
    if not stls:
        return
    click.echo(click.style("\n=== Styles ===", bold=True))
    for s in stls:
        click.echo(f"  {s['style_id']}: {s['name']} ({s['type']})")


def _pprint_headers_footers(hf_data):
    if not hf_data:
        return
    click.echo(click.style("\n=== Headers/Footers ===", bold=True))
    for hf in hf_data:
        click.echo(f"  Section[{hf['section']}] {hf['type']}: {hf['text']}")


def _pprint_images(imgs):
    if not imgs:
        return
    click.echo(click.style("\n=== Images ===", bold=True))
    for img in imgs:
        click.echo(f"  {img}")


def _pprint_properties(props):
    click.echo(click.style("\n=== Properties ===", bold=True))
    for k, v in props.items():
        if v:
            click.echo(f"  {k}: {v}")


# ---------------------------------------------------------------------------
# create
# ---------------------------------------------------------------------------

@cli.command()
@click.argument("file")
@click.option("--template", help="Use existing .docx as template")
def create_cmd(file, template):
    """Create a new .docx file."""
    if template:
        doc = Document(template)
    else:
        doc = Document()
    doc.save(file)
    click.echo(f"Created: {file}")


# ---------------------------------------------------------------------------
# paragraph
# ---------------------------------------------------------------------------

@cli.group()
def paragraph():
    """Add, list, update, or delete paragraphs."""


@paragraph.command("add")
@click.argument("file")
@click.option("--text", required=True, help="Paragraph text")
@click.option("--style", help="Paragraph style name (e.g., 'Heading 1')")
@click.option("--after", type=int, help="Insert after paragraph index (default: append to end)")
@click.option("--alignment", type=click.Choice(["left", "center", "right", "justify"]))
@click.option("-o", "--output", help="Output file path (default: overwrite input)")
def paragraph_add(file, text, style, after, alignment, output):
    """Add a paragraph to the document."""
    target, out_path, _ = _resolve_output(file, output)
    doc = _load(target)

    if after is not None and after < len(doc.paragraphs):
        # insert after specific paragraph
        ref_para = doc.paragraphs[after]
        new_para = ref_para.insert_paragraph_before(text, style)
        # Reorder: move ref_para's XML element after new_para's
        ref_para._element.addnext(new_para._element)
    else:
        doc.add_paragraph(text, style)

    if alignment:
        doc.paragraphs[-1].alignment = _parse_alignment(alignment)

    _save(doc, file, output)
    click.echo(f"Paragraph added to {out_path}")


@paragraph.command("list")
@click.argument("file")
@click.option("--json", "output_json", is_flag=True, help="Output as JSON")
def paragraph_list(file, output_json):
    """List all paragraphs with indices."""
    doc = _load(file)
    paras = []
    for i, p in enumerate(doc.paragraphs):
        pdata = {"index": i, "text": p.text}
        if p.style:
            pdata["style"] = p.style.name
        paras.append(pdata)

    if _json_out(paras, output_json):
        return
    _pprint_paragraphs(paras)


@paragraph.command("delete")
@click.argument("file")
@click.option("--index", type=int, required=True, help="Paragraph index to delete")
@click.option("-o", "--output", help="Output file path")
def paragraph_delete(file, index, output):
    """Delete a paragraph by index."""
    doc = _load(file)

    if index < 0 or index >= len(doc.paragraphs):
        raise click.BadParameter(f"Paragraph index {index} out of range (0-{len(doc.paragraphs)-1})")

    p = doc.paragraphs[index]
    p._element.getparent().remove(p._element)
    _save(doc, file, output)
    click.echo(f"Deleted paragraph [{index}] from {file}")


@paragraph.command("update")
@click.argument("file")
@click.option("--index", type=int, required=True, help="Paragraph index")
@click.option("--text", help="New paragraph text (replaces all runs)")
@click.option("--style", help="New paragraph style")
@click.option("--alignment", type=click.Choice(["left", "center", "right", "justify"]))
@click.option("--line-spacing", type=click.Choice(["single", "1.5", "double"]))
@click.option("--space-before", type=float, help="Space before in points")
@click.option("--space-after", type=float, help="Space after in points")
@click.option("-o", "--output", help="Output file path")
def paragraph_update(file, index, text, style, alignment, line_spacing, space_before, space_after, output):
    """Update a paragraph's content or formatting."""
    doc = _load(file)

    if index < 0 or index >= len(doc.paragraphs):
        raise click.BadParameter(f"Paragraph index {index} out of range (0-{len(doc.paragraphs)-1})")

    p = doc.paragraphs[index]

    if text is not None:
        p.clear()
        p.add_run(text)

    if style is not None:
        p.style = style

    if alignment is not None:
        p.alignment = _parse_alignment(alignment)

    pf = p.paragraph_format
    if line_spacing is not None:
        pf.line_spacing_rule = _parse_line_spacing(line_spacing)
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)

    _save(doc, file, output)
    click.echo(f"Updated paragraph [{index}] in {file}")


# ---------------------------------------------------------------------------
# run
# ---------------------------------------------------------------------------

@cli.group()
def run():
    """Add, list, or update runs (inline text spans) within a paragraph."""


FONT_OPTIONS = [
    click.option("--bold/--no-bold", default=None, help="Bold text"),
    click.option("--italic/--no-italic", default=None, help="Italic text"),
    click.option("--underline/--no-underline", default=None, help="Underlined text"),
    click.option("--size", type=float, help="Font size in points"),
    click.option("--font", "font_name", help="Font name"),
    click.option("--color", help="Font color (hex, e.g. FF0000 or #FF0000)"),
    click.option("--highlight", help="Highlight color (YELLOW, CYAN, RED, etc.)"),
    click.option("--superscript/--no-superscript", default=None),
    click.option("--subscript/--no-subscript", default=None),
    click.option("--strike/--no-strike", default=None, help="Strikethrough"),
    click.option("--double-strike/--no-double-strike", default=None),
    click.option("--all-caps/--no-all-caps", default=None),
    click.option("--small-caps/--no-small-caps", default=None),
    click.option("--hidden/--no-hidden", default=None),
]


def _apply_font_options(f):
    """Decorator that adds all font options to a command."""
    for opt in reversed(FONT_OPTIONS):
        f = opt(f)
    return f


@run.command("add")
@click.argument("file")
@click.option("--para", "para_index", type=int, required=True, help="Paragraph index")
@click.option("--text", required=True, help="Run text")
@click.option("--style", help="Character style name")
@click.option("-o", "--output", help="Output file path")
@_apply_font_options
def run_add(file, para_index, text, style, output, **font_kwargs):
    """Add a run to a paragraph."""
    doc = _load(file)

    if para_index < 0 or para_index >= len(doc.paragraphs):
        raise click.BadParameter(f"Paragraph index {para_index} out of range")

    p = doc.paragraphs[para_index]
    new_run = p.add_run(text, style)
    _set_run_font(new_run, **font_kwargs)

    _save(doc, file, output)
    click.echo(f"Run added to paragraph [{para_index}] in {file}")


@run.command("list")
@click.argument("file")
@click.option("--para", "para_index", type=int, required=True, help="Paragraph index")
@click.option("--json", "output_json", is_flag=True, help="Output as JSON")
def run_list(file, para_index, output_json):
    """List all runs in a paragraph."""
    doc = _load(file)

    if para_index < 0 or para_index >= len(doc.paragraphs):
        raise click.BadParameter(f"Paragraph index {para_index} out of range")

    p = doc.paragraphs[para_index]
    runs = []
    for i, r in enumerate(p.runs):
        rdata = {"index": i, "text": r.text}
        finfo = _format_font(r.font)
        if finfo:
            rdata["font"] = finfo
        if r.style:
            rdata["style"] = r.style.name
        runs.append(rdata)

    if _json_out(runs, output_json):
        return
    click.echo(f"Paragraph [{para_index}]: {p.text[:80]}")
    for r in runs:
        finfo = ""
        if r.get("font"):
            parts = []
            f = r["font"]
            if f.get("bold"):
                parts.append("bold")
            if f.get("italic"):
                parts.append("italic")
            if f.get("name"):
                parts.append(f["name"])
            if f.get("size"):
                parts.append(f"{f['size']}pt")
            if f.get("color"):
                parts.append(f"#{f['color']}")
            if parts:
                finfo = f" ({', '.join(parts)})"
        click.echo(f"  [{r['index']}]{finfo}: {r['text']}")


@run.command("update")
@click.argument("file")
@click.option("--para", "para_index", type=int, required=True, help="Paragraph index")
@click.option("--run", "run_index", type=int, required=True, help="Run index within paragraph")
@click.option("--text", help="New text for the run")
@click.option("--style", help="Character style name")
@click.option("-o", "--output", help="Output file path")
@_apply_font_options
def run_update(file, para_index, run_index, text, style, output, **font_kwargs):
    """Update a run's text or formatting."""
    doc = _load(file)

    if para_index < 0 or para_index >= len(doc.paragraphs):
        raise click.BadParameter(f"Paragraph index {para_index} out of range")

    p = doc.paragraphs[para_index]
    if run_index < 0 or run_index >= len(p.runs):
        raise click.BadParameter(f"Run index {run_index} out of range (0-{len(p.runs)-1})")

    r = p.runs[run_index]

    if text is not None:
        r.text = text

    if style is not None:
        r.style = style

    _set_run_font(r, **font_kwargs)

    _save(doc, file, output)
    click.echo(f"Updated run [{run_index}] in paragraph [{para_index}] in {file}")


# ---------------------------------------------------------------------------
# table
# ---------------------------------------------------------------------------

@cli.group()
def table():
    """Add, list, modify, or delete tables."""


@table.command("add")
@click.argument("file")
@click.option("--rows", type=int, required=True, help="Number of rows")
@click.option("--cols", type=int, required=True, help="Number of columns")
@click.option("--after", type=int, help="Insert after paragraph index")
@click.option("--style", help="Table style name")
@click.option("--header", help="Comma-separated header row text")
@click.option("--data", callback=_parse_json, help='JSON array of arrays: [["a","b"],["c","d"]]')
@click.option("-o", "--output", help="Output file path")
def table_add(file, rows, cols, after, style, header, data, output):
    """Add a table to the document."""
    doc = _load(file)

    table = doc.add_table(rows=rows, cols=cols, style=style)

    # Fill header
    if header:
        headers = [h.strip() for h in header.split(",")]
        for ci, h in enumerate(headers):
            if ci < cols:
                table.cell(0, ci).text = h

    # Fill data
    if data:
        start_row = 1 if header else 0
        for ri, row_data in enumerate(data):
            target_row = start_row + ri
            if target_row >= rows:
                break
            for ci, val in enumerate(row_data):
                if ci < cols:
                    table.cell(target_row, ci).text = str(val)

    # Move table after specific paragraph if requested
    if after is not None and after < len(doc.paragraphs):
        ref_para = doc.paragraphs[after]
        ref_para._element.addnext(table._tbl)

    _save(doc, file, output)
    click.echo(f"Table ({rows}x{cols}) added to {file}")


@table.command("list")
@click.argument("file")
@click.option("--json", "output_json", is_flag=True, help="Output as JSON")
def table_list(file, output_json):
    """List all tables with cell contents."""
    doc = _load(file)
    tbls = []
    for i, t in enumerate(doc.tables):
        rows_data = []
        for ri, row in enumerate(t.rows):
            cells = [{"col": ci, "text": cell.text} for ci, cell in enumerate(row.cells)]
            rows_data.append({"index": ri, "cells": cells})
        tdata = {
            "index": i,
            "row_count": len(t.rows),
            "col_count": len(t.columns),
            "rows": rows_data,
        }
        if t.style:
            tdata["style"] = t.style.name
        tbls.append(tdata)

    if _json_out(tbls, output_json):
        return
    _pprint_tables(tbls)


@table.command("delete")
@click.argument("file")
@click.option("--index", type=int, required=True, help="Table index to delete")
@click.option("-o", "--output", help="Output file path")
def table_delete(file, index, output):
    """Delete a table by index."""
    doc = _load(file)

    if index < 0 or index >= len(doc.tables):
        raise click.BadParameter(f"Table index {index} out of range (0-{len(doc.tables)-1})")

    t = doc.tables[index]
    t._tbl.getparent().remove(t._tbl)
    _save(doc, file, output)
    click.echo(f"Deleted table [{index}] from {file}")


@table.command("cell")
@click.argument("file")
@click.option("--table", "table_index", type=int, required=True, help="Table index")
@click.option("--row", type=int, required=True, help="Row index")
@click.option("--col", type=int, required=True, help="Column index")
@click.option("--text", help="Set cell text (omit to read)")
@click.option("-o", "--output", help="Output file path")
def table_cell(file, table_index, row, col, text, output):
    """Read or write a table cell."""
    doc = _load(file)

    if table_index < 0 or table_index >= len(doc.tables):
        raise click.BadParameter(f"Table index {table_index} out of range")

    t = doc.tables[table_index]
    if row < 0 or row >= len(t.rows):
        raise click.BadParameter(f"Row {row} out of range (0-{len(t.rows)-1})")
    if col < 0 or col >= len(t.columns):
        raise click.BadParameter(f"Col {col} out of range (0-{len(t.columns)-1})")

    cell = t.cell(row, col)

    if text is not None:
        cell.text = text
        _save(doc, file, output)
        click.echo(f"Cell [{row},{col}] updated in table [{table_index}]")
    else:
        click.echo(cell.text)


@table.command("rows")
@click.argument("file")
@click.option("--table", "table_index", type=int, required=True, help="Table index")
@click.option("--add", "add_n", type=int, help="Add N rows")
@click.option("--delete", "delete_n", type=int, help="Delete row at index")
@click.option("-o", "--output", help="Output file path")
def table_rows(file, table_index, add_n, delete_n, output):
    """Add or delete rows in a table."""
    doc = _load(file)

    if table_index < 0 or table_index >= len(doc.tables):
        raise click.BadParameter(f"Table index {table_index} out of range")

    t = doc.tables[table_index]

    if delete_n is not None:
        if delete_n < 0 or delete_n >= len(t.rows):
            raise click.BadParameter(f"Row {delete_n} out of range (0-{len(t.rows)-1})")
        tr = t.rows[delete_n]._tr
        tr.getparent().remove(tr)

    if add_n is not None:
        for _ in range(add_n):
            t.add_row()

    _save(doc, file, output)
    click.echo(f"Table [{table_index}] now has {len(t.rows)} rows")


@table.command("cols")
@click.argument("file")
@click.option("--table", "table_index", type=int, required=True, help="Table index")
@click.option("--add", "add_n", type=int, help="Add N columns")
@click.option("--delete", "delete_n", type=int, help="Delete column at index")
@click.option("-o", "--output", help="Output file path")
def table_cols(file, table_index, add_n, delete_n, output):
    """Add or delete columns in a table."""
    doc = _load(file)

    if table_index < 0 or table_index >= len(doc.tables):
        raise click.BadParameter(f"Table index {table_index} out of range")

    t = doc.tables[table_index]

    if delete_n is not None:
        if delete_n < 0 or delete_n >= len(t.columns):
            raise click.BadParameter(f"Column {delete_n} out of range (0-{len(t.columns)-1})")
        col = t.columns[delete_n]
        for cell in col.cells:
            cell._tc.getparent().remove(cell._tc)

    if add_n is not None:
        for _ in range(add_n):
            t.add_column(Inches(1.0))

    _save(doc, file, output)
    click.echo(f"Table [{table_index}] now has {len(t.columns)} columns")


# ---------------------------------------------------------------------------
# image
# ---------------------------------------------------------------------------

@cli.group()
def image():
    """Add or list images."""


@image.command("add")
@click.argument("file")
@click.argument("image_path")
@click.option("--para", type=int, help="Paragraph index to add image in a new run")
@click.option("--width", type=float, help="Width in inches")
@click.option("--height", type=float, help="Height in inches")
@click.option("-o", "--output", help="Output file path")
def image_add(file, image_path, para, width, height, output):
    """Add an image to the document."""
    doc = _load(file)
    img_path = Path(image_path)
    if not img_path.exists():
        raise click.BadParameter(f"Image file not found: {image_path}")

    w = Inches(width) if width else None
    h = Inches(height) if height else None

    if para is not None:
        if para < 0 or para >= len(doc.paragraphs):
            raise click.BadParameter(f"Paragraph index {para} out of range")
        run = doc.paragraphs[para].add_run()
        run.add_picture(str(img_path), width=w, height=h)
    else:
        doc.add_picture(str(img_path), width=w, height=h)

    _save(doc, file, output)
    click.echo(f"Image '{image_path}' added to {file}")


@image.command("list")
@click.argument("file")
@click.option("--json", "output_json", is_flag=True, help="Output as JSON")
def image_list(file, output_json):
    """List images in the document."""
    doc = _load(file)
    imgs = []

    # Inline shapes
    for s in doc.inline_shapes:
        imgs.append({"type": str(s.type), "width": s.width, "height": s.height})

    # Images in runs
    for i, p in enumerate(doc.paragraphs):
        for j, r in enumerate(p.runs):
            drawings = r._element.findall(qn("w:drawing"))
            for d in drawings:
                imgs.append({"paragraph": i, "run": j, "type": "drawing"})

    if _json_out(imgs, output_json):
        return
    _pprint_images(imgs)


# ---------------------------------------------------------------------------
# section
# ---------------------------------------------------------------------------

@cli.group()
def section():
    """List, add, or update document sections."""


@section.command("list")
@click.argument("file")
@click.option("--json", "output_json", is_flag=True, help="Output as JSON")
def section_list(file, output_json):
    """List all sections."""
    doc = _load(file)
    secs = []
    for i, s in enumerate(doc.sections):
        sdata = {
            "index": i,
            "page_width": s.page_width.pt if s.page_width else None,
            "page_height": s.page_height.pt if s.page_height else None,
            "orientation": str(s.orientation),
            "start_type": str(s.start_type),
            "left_margin": s.left_margin.pt if s.left_margin else None,
            "right_margin": s.right_margin.pt if s.right_margin else None,
            "top_margin": s.top_margin.pt if s.top_margin else None,
            "bottom_margin": s.bottom_margin.pt if s.bottom_margin else None,
        }
        secs.append(sdata)

    if _json_out(secs, output_json):
        return
    _pprint_sections(secs)


@section.command("add")
@click.argument("file")
@click.option("--type", "start_type", type=click.Choice(["continuous", "new_page", "new_column", "even_page", "odd_page"]),
              help="Section break type")
@click.option("--orientation", type=click.Choice(["portrait", "landscape"]))
@click.option("-o", "--output", help="Output file path")
def section_add(file, start_type, orientation, output):
    """Add a section break."""
    doc = _load(file)

    st = {
        "continuous": WD_SECTION_START.CONTINUOUS,
        "new_page": WD_SECTION_START.NEW_PAGE,
        "new_column": WD_SECTION_START.NEW_COLUMN,
        "even_page": WD_SECTION_START.EVEN_PAGE,
        "odd_page": WD_SECTION_START.ODD_PAGE,
    }
    new_section = doc.add_section(st.get(start_type, WD_SECTION_START.NEW_PAGE))

    if orientation == "landscape":
        new_section.orientation = WD_ORIENT.LANDSCAPE
    elif orientation == "portrait":
        new_section.orientation = WD_ORIENT.PORTRAIT

    _save(doc, file, output)
    click.echo(f"Section added to {file} (total: {len(doc.sections)})")


@section.command("update")
@click.argument("file")
@click.option("--index", type=int, required=True, help="Section index")
@click.option("--orientation", type=click.Choice(["portrait", "landscape"]))
@click.option("--margin-top", type=float, help="Top margin in points")
@click.option("--margin-bottom", type=float, help="Bottom margin in points")
@click.option("--margin-left", type=float, help="Left margin in points")
@click.option("--margin-right", type=float, help="Right margin in points")
@click.option("--page-width", type=float, help="Page width in points")
@click.option("--page-height", type=float, help="Page height in points")
@click.option("-o", "--output", help="Output file path")
def section_update(file, index, orientation, margin_top, margin_bottom,
                   margin_left, margin_right, page_width, page_height, output):
    """Update section properties."""
    doc = _load(file)

    if index < 0 or index >= len(doc.sections):
        raise click.BadParameter(f"Section index {index} out of range (0-{len(doc.sections)-1})")

    s = doc.sections[index]

    if orientation == "landscape":
        s.orientation = WD_ORIENT.LANDSCAPE
    elif orientation == "portrait":
        s.orientation = WD_ORIENT.PORTRAIT

    if margin_top is not None:
        s.top_margin = Pt(margin_top)
    if margin_bottom is not None:
        s.bottom_margin = Pt(margin_bottom)
    if margin_left is not None:
        s.left_margin = Pt(margin_left)
    if margin_right is not None:
        s.right_margin = Pt(margin_right)
    if page_width is not None:
        s.page_width = Pt(page_width)
    if page_height is not None:
        s.page_height = Pt(page_height)

    _save(doc, file, output)
    click.echo(f"Updated section [{index}] in {file}")


# ---------------------------------------------------------------------------
# header-footer
# ---------------------------------------------------------------------------

@cli.group(name="header-footer")
def header_footer():
    """Add or list headers and footers."""


@header_footer.command("list")
@click.argument("file")
@click.option("--section", "section_index", type=int, default=0, help="Section index")
@click.option("--json", "output_json", is_flag=True, help="Output as JSON")
def hf_list(file, section_index, output_json):
    """List header/footer content."""
    doc = _load(file)

    if section_index < 0 or section_index >= len(doc.sections):
        raise click.BadParameter(f"Section index {section_index} out of range")

    s = doc.sections[section_index]
    hf_data = []
    for hf_type in ["header", "footer", "first_page_header", "first_page_footer",
                     "even_page_header", "even_page_footer"]:
        try:
            hf = getattr(s, hf_type, None)
            if hf:
                paras = [p.text for p in hf.paragraphs]
                hf_data.append({"type": hf_type, "paragraphs": paras})
        except Exception:
            pass

    if _json_out(hf_data, output_json):
        return
    for hf in hf_data:
        click.echo(f"{hf['type']}:")
        for p in hf["paragraphs"]:
            click.echo(f"  {p}")


@header_footer.command("add")
@click.argument("file")
@click.option("--type", "hf_type", type=click.Choice(["header", "footer", "first_page_header", "first_page_footer"]),
              required=True, help="Header/footer type")
@click.option("--text", required=True, help="Text to add")
@click.option("--section", "section_index", type=int, default=0, help="Section index")
@click.option("--style", help="Paragraph style")
@click.option("--alignment", type=click.Choice(["left", "center", "right", "justify"]))
@click.option("-o", "--output", help="Output file path")
def hf_add(file, hf_type, text, section_index, style, alignment, output):
    """Add text to a header or footer."""
    doc = _load(file)

    if section_index < 0 or section_index >= len(doc.sections):
        raise click.BadParameter(f"Section index {section_index} out of range")

    s = doc.sections[section_index]

    # Enable different first page if needed
    if hf_type.startswith("first_page"):
        s.different_first_page_header_footer = True

    hf_obj = getattr(s, hf_type)
    p = hf_obj.add_paragraph(text, style)
    if alignment:
        p.alignment = _parse_alignment(alignment)

    _save(doc, file, output)
    click.echo(f"Added '{hf_type}' text to section [{section_index}] in {file}")


# ---------------------------------------------------------------------------
# style
# ---------------------------------------------------------------------------

@cli.group()
def style():
    """List document styles."""


@style.command("list")
@click.argument("file")
@click.option("--type", "style_type", type=click.Choice(["paragraph", "character", "table", "list"]),
              help="Filter by style type")
@click.option("--json", "output_json", is_flag=True, help="Output as JSON")
def style_list(file, style_type, output_json):
    """List available styles in the document."""
    doc = _load(file)
    stls = []
    for s in doc.styles:
        st = str(s.type)
        if style_type and st != style_type.upper():
            # Allow "character" to match "CHARACTER (2)"
            if style_type.upper() not in st:
                continue
        stls.append({"name": s.name, "style_id": s.style_id, "type": st})

    if _json_out(stls, output_json):
        return
    _pprint_styles(stls)


# ---------------------------------------------------------------------------
# properties
# ---------------------------------------------------------------------------

@cli.group()
def properties():
    """Get or set document properties."""


@properties.command("get")
@click.argument("file")
@click.option("--json", "output_json", is_flag=True, help="Output as JSON")
def props_get(file, output_json):
    """Get document core properties."""
    doc = _load(file)
    cp = doc.core_properties
    props = {
        "author": cp.author,
        "title": cp.title,
        "subject": cp.subject,
        "keywords": cp.keywords,
        "category": cp.category,
        "comments": cp.comments,
        "created": str(cp.created) if cp.created else None,
        "modified": str(cp.modified) if cp.modified else None,
        "last_modified_by": cp.last_modified_by,
        "revision": cp.revision,
    }

    if _json_out(props, output_json):
        return
    for k, v in props.items():
        if v:
            click.echo(f"{k}: {v}")


@properties.command("set")
@click.argument("file")
@click.option("--title")
@click.option("--author")
@click.option("--subject")
@click.option("--keywords")
@click.option("--category")
@click.option("--comments")
@click.option("-o", "--output", help="Output file path")
def props_set(file, title, author, subject, keywords, category, comments, output):
    """Set document core properties."""
    doc = _load(file)
    cp = doc.core_properties

    if title is not None:
        cp.title = title
    if author is not None:
        cp.author = author
    if subject is not None:
        cp.subject = subject
    if keywords is not None:
        cp.keywords = keywords
    if category is not None:
        cp.category = category
    if comments is not None:
        cp.comments = comments

    _save(doc, file, output)
    click.echo(f"Properties updated in {file}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    cli()
